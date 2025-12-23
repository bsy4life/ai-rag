# personal_kb.py - 個人知識庫完整實作（生產級）
"""
功能：
1. 多格式文件解析（DOCX, PDF, TXT, MD, XLSX）
2. 圖片提取與管理
3. 關鍵字索引
4. 個人向量庫
5. 混合檢索
6. 與主系統整合
"""

import os
import re
import json
import hashlib
import zipfile
import shutil
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, asdict, field
from threading import Lock

# 嘗試從 config 導入，如果失敗則使用預設值
try:
    from config import (
        PERSONAL_KB_DIR, PERSONAL_KB_CONFIG, EMBEDDING_MODEL,
        CHUNK_CONFIGS, KEYWORD_PATTERNS, CHINESE_STOPWORDS,
    )
except ImportError:
    # 預設值
    PERSONAL_KB_DIR = os.getenv("PERSONAL_KB_DIR", "/app/data/personal_kb")
    EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "text-embedding-3-small")
    
    @dataclass
    class _PersonalKBConfig:
        enabled: bool = True
        max_file_size_mb: int = 50
        allowed_extensions: List[str] = field(default_factory=lambda: [
            ".docx", ".pdf", ".txt", ".md", ".xlsx", ".csv", ".png", ".jpg", ".jpeg"
        ])
        auto_extract_images: bool = True
        max_images_per_doc: int = 50
    
    PERSONAL_KB_CONFIG = _PersonalKBConfig()
    
    CHUNK_CONFIGS = {
        "personal": type('obj', (object,), {
            'chunk_size': 500,
            'chunk_overlap': 80,
        })()
    }
    
    KEYWORD_PATTERNS = [
        r'[A-Z]{2,}\d*[A-Za-z]*',
        r'[A-Za-z]+[-_][A-Za-z0-9]+',
        r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}',
        r'[A-Za-z]+\.[A-Za-z]{2,4}',
        r'No\.\s*\d+',
    ]
    
    CHINESE_STOPWORDS = {
        '的', '是', '在', '和', '了', '有', '這', '個', '不', '為',
        '上', '下', '中', '請', '到', '把', '被', '讓', '給', '跟',
    }

logger = logging.getLogger(__name__)

# ═══════════════════════════════════════════════════════════════
# 資料結構
# ═══════════════════════════════════════════════════════════════

@dataclass
class ExtractedImage:
    """提取的圖片"""
    name: str
    path: str
    size: int = 0
    width: int = 0
    height: int = 0
    paragraph_idx: Optional[int] = None
    paragraph_text: Optional[str] = None
    
@dataclass
class DocumentChunk:
    """文件區塊"""
    content: str
    chunk_idx: int
    doc_id: str
    metadata: Dict = field(default_factory=dict)

@dataclass
class ProcessedDocument:
    """處理後的文件"""
    doc_id: str
    filename: str
    original_path: str
    file_type: str
    file_size: int
    text: str
    chunks: List[DocumentChunk]
    images: List[ExtractedImage]
    keywords: List[str]
    processed_at: str
    status: str = "processed"
    error: Optional[str] = None

@dataclass
class SearchResult:
    """搜尋結果"""
    doc_id: str
    filename: str
    content: str
    score: float
    match_type: str  # keyword, vector, hybrid
    images: List[Dict] = field(default_factory=list)
    metadata: Dict = field(default_factory=dict)

# ═══════════════════════════════════════════════════════════════
# 文件解析器
# ═══════════════════════════════════════════════════════════════

class DocumentParser:
    """多格式文件解析器"""
    
    @staticmethod
    def parse(file_path: str, output_dir: str) -> Tuple[str, List[ExtractedImage]]:
        """
        解析文件
        
        Returns:
            (文字內容, 圖片列表)
        """
        ext = Path(file_path).suffix.lower()
        
        parsers = {
            '.docx': DocumentParser._parse_docx,
            '.pdf': DocumentParser._parse_pdf,
            '.txt': DocumentParser._parse_text,
            '.md': DocumentParser._parse_text,
            '.xlsx': DocumentParser._parse_xlsx,
            '.csv': DocumentParser._parse_csv,
            '.png': DocumentParser._parse_image,
            '.jpg': DocumentParser._parse_image,
            '.jpeg': DocumentParser._parse_image,
            '.gif': DocumentParser._parse_image,
        }
        
        parser = parsers.get(ext)
        if not parser:
            raise ValueError(f"不支援的檔案格式: {ext}")
        
        return parser(file_path, output_dir)
    
    @staticmethod
    def _parse_docx(file_path: str, output_dir: str) -> Tuple[str, List[ExtractedImage]]:
        """解析 DOCX"""
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError("需要安裝 python-docx: pip install python-docx")
        
        doc = Document(file_path)
        images_dir = os.path.join(output_dir, "images")
        os.makedirs(images_dir, exist_ok=True)
        
        # 提取圖片
        images = []
        with zipfile.ZipFile(file_path, 'r') as z:
            for name in z.namelist():
                if name.startswith('word/media/'):
                    img_data = z.read(name)
                    ext = name.split('.')[-1].lower()
                    if ext not in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp']:
                        continue
                    
                    img_name = f"img_{len(images)+1:03d}.{ext}"
                    img_path = os.path.join(images_dir, img_name)
                    
                    with open(img_path, 'wb') as f:
                        f.write(img_data)
                    
                    images.append(ExtractedImage(
                        name=img_name,
                        path=img_path,
                        size=len(img_data),
                    ))
        
        # 提取文字並關聯圖片
        text_parts = []
        para_idx = 0
        image_idx = 0
        
        for para in doc.paragraphs:
            para_text = para.text.strip()
            
            # 檢查段落中的圖片
            drawings = para._element.findall('.//' + qn('w:drawing'))
            if drawings and image_idx < len(images):
                for _ in drawings:
                    if image_idx < len(images):
                        images[image_idx].paragraph_idx = para_idx
                        images[image_idx].paragraph_text = para_text[:100] if para_text else f"步驟 {para_idx + 1}"
                        text_parts.append(f"\n[📷 圖片: {images[image_idx].name}]\n")
                        image_idx += 1
            
            if para_text:
                text_parts.append(para_text)
            
            para_idx += 1
        
        return "\n".join(text_parts), images
    
    @staticmethod
    def _parse_pdf(file_path: str, output_dir: str) -> Tuple[str, List[ExtractedImage]]:
        """解析 PDF"""
        images_dir = os.path.join(output_dir, "images")
        os.makedirs(images_dir, exist_ok=True)
        
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            text_parts = []
            images = []
            
            for page_num, page in enumerate(doc):
                # 提取文字
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"[📄 第 {page_num + 1} 頁]\n{text}")
                
                # 提取圖片
                for img_idx, img in enumerate(page.get_images()):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        
                        img_name = f"img_p{page_num+1}_{img_idx+1:03d}.png"
                        img_path = os.path.join(images_dir, img_name)
                        
                        if pix.n < 5:
                            pix.save(img_path)
                        else:
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                            pix.save(img_path)
                        
                        images.append(ExtractedImage(
                            name=img_name,
                            path=img_path,
                            size=os.path.getsize(img_path),
                            paragraph_idx=page_num,
                            paragraph_text=f"第 {page_num + 1} 頁",
                        ))
                        
                        text_parts.append(f"\n[📷 圖片: {img_name}]\n")
                    except:
                        pass
            
            doc.close()
            return "\n".join(text_parts), images
            
        except ImportError:
            # 備用：pdfplumber
            try:
                import pdfplumber
                
                text_parts = []
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text:
                            text_parts.append(f"[📄 第 {page_num + 1} 頁]\n{text}")
                
                return "\n".join(text_parts), []
                
            except ImportError:
                raise ImportError("需要安裝 PyMuPDF 或 pdfplumber")
    
    @staticmethod
    def _parse_text(file_path: str, output_dir: str) -> Tuple[str, List[ExtractedImage]]:
        """解析純文字"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        return text, []
    
    @staticmethod
    def _parse_xlsx(file_path: str, output_dir: str) -> Tuple[str, List[ExtractedImage]]:
        """解析 Excel"""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("需要安裝 pandas 和 openpyxl")
        
        xls = pd.ExcelFile(file_path)
        text_parts = []
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            text_parts.append(f"[📊 工作表: {sheet_name}]")
            text_parts.append(df.to_markdown(index=False))
            text_parts.append("")
        
        return "\n".join(text_parts), []
    
    @staticmethod
    def _parse_csv(file_path: str, output_dir: str) -> Tuple[str, List[ExtractedImage]]:
        """解析 CSV"""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("需要安裝 pandas")
        
        df = pd.read_csv(file_path)
        return df.to_markdown(index=False), []
    
    @staticmethod
    def _parse_image(file_path: str, output_dir: str) -> Tuple[str, List[ExtractedImage]]:
        """解析圖片檔案（使用 OCR 或作為圖片索引）"""
        import shutil
        from pathlib import Path
        
        images_dir = os.path.join(output_dir, "images")
        os.makedirs(images_dir, exist_ok=True)
        
        # 複製圖片到 extracted 目錄
        filename = Path(file_path).name
        img_path = os.path.join(images_dir, filename)
        shutil.copy2(file_path, img_path)
        
        # 取得圖片資訊
        img_size = os.path.getsize(file_path)
        width, height = 0, 0
        
        try:
            from PIL import Image
            with Image.open(file_path) as img:
                width, height = img.size
        except:
            pass
        
        # 建立圖片記錄
        image_record = ExtractedImage(
            name=filename,
            path=img_path,
            size=img_size,
            width=width,
            height=height,
            paragraph_idx=0,
            paragraph_text=f"上傳的圖片: {filename}"
        )
        
        # 嘗試 OCR 提取文字
        ocr_text = ""
        
        # 方法 1: 嘗試 pytesseract（如果有安裝）
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(file_path)
            # 支援中文+英文
            ocr_text = pytesseract.image_to_string(img, lang='chi_tra+eng')
            if ocr_text.strip():
                logger.info(f"✅ OCR 成功提取文字: {len(ocr_text)} 字元")
        except ImportError:
            logger.debug("pytesseract 未安裝，跳過 OCR")
        except Exception as e:
            logger.debug(f"OCR 失敗: {e}")
        
        # 方法 2: 如果 OCR 失敗，使用檔名作為描述
        if not ocr_text.strip():
            ocr_text = f"[圖片檔案]\n檔名: {filename}\n尺寸: {width}x{height}\n大小: {img_size} bytes"
        
        # 組合文字內容
        text_content = f"""# 圖片: {filename}

{ocr_text}

[📷 圖片: {filename}]
"""
        
        return text_content, [image_record]

# ═══════════════════════════════════════════════════════════════
# 關鍵字索引
# ═══════════════════════════════════════════════════════════════

class PersonalKeywordIndex:
    """個人知識庫關鍵字索引"""
    
    def __init__(self, index_path: str):
        self.index_path = index_path
        self.index: Dict[str, List[Tuple[str, float]]] = {}
        self.doc_keywords: Dict[str, List[str]] = {}
        self._lock = Lock()
        self._load()
    
    def _extract_keywords(self, text: str) -> List[str]:
        """提取關鍵字"""
        keywords = set()
        
        for pattern in KEYWORD_PATTERNS:
            matches = re.findall(pattern, text, re.IGNORECASE)
            keywords.update(matches)
        
        # 中文詞彙
        chinese = re.findall(r'[\u4e00-\u9fa5]{2,8}', text)
        keywords.update(w for w in chinese if w not in CHINESE_STOPWORDS and len(w) >= 2)
        
        return list(keywords)
    
    def add(self, doc_id: str, text: str):
        """添加文件"""
        keywords = self._extract_keywords(text)
        
        with self._lock:
            self.doc_keywords[doc_id] = keywords
            
            for kw in keywords:
                kw_lower = kw.lower()
                if kw_lower not in self.index:
                    self.index[kw_lower] = []
                
                if not any(d == doc_id for d, _ in self.index[kw_lower]):
                    self.index[kw_lower].append((doc_id, 1.0))
        
        self._save()
        return keywords
    
    def remove(self, doc_id: str):
        """移除文件"""
        with self._lock:
            if doc_id in self.doc_keywords:
                for kw in self.doc_keywords[doc_id]:
                    kw_lower = kw.lower()
                    if kw_lower in self.index:
                        self.index[kw_lower] = [(d, s) for d, s in self.index[kw_lower] if d != doc_id]
                del self.doc_keywords[doc_id]
        
        self._save()
    
    def search(self, query: str, top_k: int = 10) -> List[Tuple[str, float]]:
        """搜尋"""
        query_keywords = self._extract_keywords(query)
        query_words = set(query.lower().split())
        all_terms = set(kw.lower() for kw in query_keywords) | query_words
        
        doc_scores: Dict[str, float] = {}
        
        for term in all_terms:
            if term in self.index:
                for doc_id, weight in self.index[term]:
                    doc_scores[doc_id] = doc_scores.get(doc_id, 0) + 2 * weight
            
            for indexed_kw, doc_list in self.index.items():
                if len(term) >= 2 and (term in indexed_kw or indexed_kw in term):
                    for doc_id, weight in doc_list:
                        doc_scores[doc_id] = doc_scores.get(doc_id, 0) + weight
        
        results = sorted(doc_scores.items(), key=lambda x: x[1], reverse=True)
        return results[:top_k]
    
    def _save(self):
        """儲存"""
        data = {
            "index": {k: list(v) for k, v in self.index.items()},
            "doc_keywords": self.doc_keywords,
        }
        with open(self.index_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    
    def _load(self):
        """載入"""
        if os.path.exists(self.index_path):
            try:
                with open(self.index_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.index = {k: [tuple(x) for x in v] for k, v in data.get("index", {}).items()}
                    self.doc_keywords = data.get("doc_keywords", {})
            except:
                pass

# ═══════════════════════════════════════════════════════════════
# 個人知識庫
# ═══════════════════════════════════════════════════════════════

class PersonalKnowledgeBase:
    """個人知識庫"""
    
    def __init__(self, user_id: str):
        self.user_id = user_id
        self.base_dir = os.path.join(PERSONAL_KB_DIR, user_id)
        self.docs_dir = os.path.join(self.base_dir, "documents")
        self.extracted_dir = os.path.join(self.base_dir, "extracted")
        self.vectordb_dir = os.path.join(self.base_dir, "vectordb")
        self.metadata_path = os.path.join(self.base_dir, "metadata.json")
        
        # 建立目錄
        for d in [self.docs_dir, self.extracted_dir, self.vectordb_dir]:
            os.makedirs(d, exist_ok=True)
        
        # 載入
        self.metadata = self._load_metadata()
        self.keyword_index = PersonalKeywordIndex(
            os.path.join(self.base_dir, "keyword_index.json")
        )
        
        self._vectordb = None
        self._lock = Lock()
    
    def _load_metadata(self) -> Dict:
        """載入 metadata"""
        if os.path.exists(self.metadata_path):
            with open(self.metadata_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        
        return {
            "user_id": self.user_id,
            "documents": {},
            "stats": {
                "total_documents": 0,
                "total_chunks": 0,
                "total_images": 0,
                "last_updated": None,
            }
        }
    
    def _save_metadata(self):
        """儲存 metadata"""
        self.metadata["stats"]["last_updated"] = datetime.now().isoformat()
        with open(self.metadata_path, 'w', encoding='utf-8') as f:
            json.dump(self.metadata, f, ensure_ascii=False, indent=2)
    
    def _get_vectordb(self):
        """取得向量庫"""
        if self._vectordb is None:
            try:
                from langchain_chroma import Chroma
                from langchain_openai import OpenAIEmbeddings
                import chromadb
                
                logger.info(f"🔧 初始化個人向量庫: {self.vectordb_dir}")
                
                embedding = OpenAIEmbeddings(model=EMBEDDING_MODEL)
                
                # 確保目錄存在
                os.makedirs(self.vectordb_dir, exist_ok=True)
                
                client = chromadb.PersistentClient(path=self.vectordb_dir)
                
                self._vectordb = Chroma(
                    client=client,
                    collection_name=f"personal_{self.user_id}",
                    embedding_function=embedding
                )
                logger.info(f"✅ 個人向量庫初始化成功: personal_{self.user_id}")
            except Exception as e:
                logger.error(f"❌ 向量庫初始化失敗: {e}")
                import traceback
                logger.error(traceback.format_exc())
        
        return self._vectordb
    
    def add_document(self, file_path: str, filename: str = None) -> Dict:
        """添加文件"""
        filename = filename or os.path.basename(file_path)
        
        # 檢查檔案大小
        file_size = os.path.getsize(file_path)
        max_size = PERSONAL_KB_CONFIG.max_file_size_mb * 1024 * 1024
        if file_size > max_size:
            return {
                "success": False,
                "error": f"檔案太大（{file_size/1024/1024:.1f}MB > {PERSONAL_KB_CONFIG.max_file_size_mb}MB）"
            }
        
        # 檢查格式
        ext = Path(filename).suffix.lower()
        if ext not in PERSONAL_KB_CONFIG.allowed_extensions:
            return {
                "success": False,
                "error": f"不支援的格式: {ext}"
            }
        
        # 生成 ID
        with open(file_path, 'rb') as f:
            file_hash = hashlib.md5(f.read()).hexdigest()[:12]
        doc_id = f"doc_{file_hash}"
        
        # 檢查重複
        if doc_id in self.metadata["documents"]:
            return {
                "success": False,
                "error": "文件已存在",
                "doc_id": doc_id,
            }
        
        # 建立輸出目錄
        doc_output_dir = os.path.join(self.extracted_dir, doc_id)
        os.makedirs(doc_output_dir, exist_ok=True)
        
        try:
            # 解析文件
            text, images = DocumentParser.parse(file_path, doc_output_dir)
            
            # 儲存原始文字
            text_path = os.path.join(doc_output_dir, f"{doc_id}.md")
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(f"# {filename}\n\n{text}")
            
            # 切塊
            chunks = self._split_text(text, doc_id, filename)
            
            # 添加到向量庫
            vectordb = self._get_vectordb()
            if vectordb and chunks:
                from langchain_core.documents import Document
                docs = [
                    Document(
                        page_content=chunk.content,
                        metadata={
                            "doc_id": doc_id,
                            "filename": filename,
                            "chunk_idx": chunk.chunk_idx,
                        }
                    )
                    for chunk in chunks
                ]
                logger.info(f"📤 開始 embedding {len(docs)} 個 chunks...")
                vectordb.add_documents(docs)
                logger.info(f"✅ Embedding 完成: {len(docs)} chunks 已加入向量庫")
            else:
                logger.warning(f"⚠️ 跳過 embedding: vectordb={vectordb is not None}, chunks={len(chunks) if chunks else 0}")
            
            # 添加到關鍵字索引
            keywords = self.keyword_index.add(doc_id, text)
            
            # 更新 metadata
            self.metadata["documents"][doc_id] = {
                "filename": filename,
                "file_type": ext,
                "file_size": file_size,
                "upload_time": datetime.now().isoformat(),
                "status": "indexed",
                "text_path": text_path,
                "chunk_count": len(chunks),
                "images": [asdict(img) for img in images],
                "keywords": keywords[:50],
            }
            
            self.metadata["stats"]["total_documents"] += 1
            self.metadata["stats"]["total_chunks"] += len(chunks)
            self.metadata["stats"]["total_images"] += len(images)
            self._save_metadata()
            
            logger.info(f"✅ 文件已添加: {filename} ({len(chunks)} chunks, {len(images)} images)")
            
            return {
                "success": True,
                "doc_id": doc_id,
                "filename": filename,
                "chunks": len(chunks),
                "images": len(images),
                "keywords": keywords[:20],
            }
            
        except Exception as e:
            logger.error(f"文件處理失敗: {e}")
            # 清理
            if os.path.exists(doc_output_dir):
                shutil.rmtree(doc_output_dir)
            
            return {
                "success": False,
                "error": str(e),
            }
    
    def _split_text(self, text: str, doc_id: str, filename: str) -> List[DocumentChunk]:
        """切割文字"""
        config = CHUNK_CONFIGS.get("personal", CHUNK_CONFIGS["technical"])
        chunk_size = config.chunk_size
        overlap = config.chunk_overlap
        
        prefix = f"【{filename}】\n"
        chunks = []
        
        # 按段落切割
        paragraphs = re.split(r'\n\n+', text)
        current = prefix
        
        for para in paragraphs:
            if len(current) + len(para) > chunk_size:
                if current.strip() and current != prefix:
                    chunks.append(DocumentChunk(
                        content=current.strip(),
                        chunk_idx=len(chunks),
                        doc_id=doc_id,
                    ))
                # 保留重疊
                overlap_text = current[-overlap:] if len(current) > overlap else ""
                current = prefix + overlap_text + para
            else:
                current += "\n\n" + para if current != prefix else para
        
        if current.strip() and current != prefix:
            chunks.append(DocumentChunk(
                content=current.strip(),
                chunk_idx=len(chunks),
                doc_id=doc_id,
            ))
        
        return chunks
    
    def remove_document(self, doc_id: str) -> bool:
        """移除文件"""
        if doc_id not in self.metadata["documents"]:
            return False
        
        doc_info = self.metadata["documents"][doc_id]
        
        # 從向量庫移除
        vectordb = self._get_vectordb()
        if vectordb:
            try:
                vectordb._collection.delete(where={"doc_id": doc_id})
            except:
                pass
        
        # 從關鍵字索引移除
        self.keyword_index.remove(doc_id)
        
        # 刪除檔案
        doc_dir = os.path.join(self.extracted_dir, doc_id)
        if os.path.exists(doc_dir):
            shutil.rmtree(doc_dir)
        
        # 更新 metadata
        self.metadata["stats"]["total_documents"] -= 1
        self.metadata["stats"]["total_chunks"] -= doc_info.get("chunk_count", 0)
        self.metadata["stats"]["total_images"] -= len(doc_info.get("images", []))
        del self.metadata["documents"][doc_id]
        self._save_metadata()
        
        logger.info(f"✅ 文件已移除: {doc_info.get('filename')}")
        return True
    
    def search(self, query: str, top_k: int = 5, include_images: bool = False) -> List[SearchResult]:
        """混合搜尋"""
        results: List[SearchResult] = []
        seen_docs = set()
        
        # 1. 關鍵字搜尋
        kw_results = self.keyword_index.search(query, top_k=top_k)
        for doc_id, score in kw_results:
            if doc_id in self.metadata["documents"] and doc_id not in seen_docs:
                doc_info = self.metadata["documents"][doc_id]
                results.append(SearchResult(
                    doc_id=doc_id,
                    filename=doc_info["filename"],
                    content=self._get_doc_preview(doc_id),
                    score=score * 2,
                    match_type="keyword",
                    images=(doc_info.get("images", [])[:3] if include_images else []),
                ))
                seen_docs.add(doc_id)
        
        # 2. 向量搜尋
        vectordb = self._get_vectordb()
        if vectordb:
            try:
                vector_docs = vectordb.similarity_search_with_score(query, k=top_k)
                for doc, score in vector_docs:
                    doc_id = doc.metadata.get("doc_id")
                    if doc_id and doc_id not in seen_docs and doc_id in self.metadata["documents"]:
                        doc_info = self.metadata["documents"][doc_id]
                        results.append(SearchResult(
                            doc_id=doc_id,
                            filename=doc_info["filename"],
                            content=doc.page_content[:300],
                            score=1.0 / (1.0 + score),
                            match_type="vector",
                            images=(doc_info.get("images", [])[:3] if include_images else []),
                        ))
                        seen_docs.add(doc_id)
            except Exception as e:
                logger.warning(f"向量搜尋失敗: {e}")
        
        # 排序
        results.sort(key=lambda x: x.score, reverse=True)
        return results[:top_k]
    
    def _get_doc_preview(self, doc_id: str, max_length: int = 300) -> str:
        """取得文件預覽"""
        doc_info = self.metadata["documents"].get(doc_id, {})
        text_path = doc_info.get("text_path")
        
        if text_path and os.path.exists(text_path):
            with open(text_path, 'r', encoding='utf-8') as f:
                text = f.read()
                return text[:max_length] + "..." if len(text) > max_length else text
        
        return ""
    
    def get_image_path(self, doc_id: str, image_name: str) -> Optional[str]:
        """取得圖片路徑"""
        path = os.path.join(self.extracted_dir, doc_id, "images", image_name)
        return path if os.path.exists(path) else None
    
    def get_document_content(self, doc_id: str) -> Optional[str]:
        """取得文件完整內容"""
        doc_info = self.metadata["documents"].get(doc_id)
        if not doc_info:
            return None
        
        text_path = doc_info.get("text_path")
        if text_path and os.path.exists(text_path):
            with open(text_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        return None
    
    def list_documents(self) -> List[Dict]:
        """列出所有文件"""
        return [
            {
                "doc_id": doc_id,
                "filename": info["filename"],
                "file_type": info["file_type"],
                "file_size": info["file_size"],
                "upload_time": info["upload_time"],
                "status": info["status"],
                "chunks": info.get("chunk_count", 0),
                "images": len(info.get("images", [])),
            }
            for doc_id, info in self.metadata["documents"].items()
        ]
    
    def get_stats(self) -> Dict:
        """取得統計"""
        return self.metadata["stats"]

# ═══════════════════════════════════════════════════════════════
# 便捷函數
# ═══════════════════════════════════════════════════════════════

_kb_cache: Dict[str, PersonalKnowledgeBase] = {}
_kb_lock = Lock()

def get_personal_kb(user_id: str) -> PersonalKnowledgeBase:
    """取得用戶的個人知識庫"""
    global _kb_cache
    
    if user_id not in _kb_cache:
        with _kb_lock:
            if user_id not in _kb_cache:
                _kb_cache[user_id] = PersonalKnowledgeBase(user_id)
    
    return _kb_cache[user_id]

def add_document(user_id: str, file_path: str, filename: str = None) -> Dict:
    """添加文件"""
    return get_personal_kb(user_id).add_document(file_path, filename)

def search_personal(user_id: str, query: str, top_k: int = 5, include_images: bool = False) -> List[SearchResult]:
    """搜尋個人知識庫"""
    return get_personal_kb(user_id).search(query, top_k, include_images=include_images)
